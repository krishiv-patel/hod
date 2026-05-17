"""Inspect a few sample documents to understand their structure."""
import os
from docx import Document

BASE_DIR = r"d:\HOD\HOD Minutes (Updated on 24 March 2026)\HOD Minutes (Updated on 24 March 2026)"

# Pick a variety of samples: early, mid, late, small, large
samples = [
    "298 HOD MoM (1st Apr'26).docx",   # recent with date in filename
    "299 HOD (8 Apr).docx",             # most recent
    "1st HOD MoM.docx",                 # first meeting
    "134th HOD MoM.docx",               # small file (~21KB, likely different format)
    "277 HOD MoM.docx",                 # small file
    "213 HOD MOM.docx",                 # different naming
]

for fname in samples:
    fpath = os.path.join(BASE_DIR, fname)
    if not os.path.exists(fpath):
        print(f"\n{'='*80}\n*** FILE NOT FOUND: {fname}\n")
        continue
    print(f"\n{'='*80}")
    print(f"FILE: {fname}  (size: {os.path.getsize(fpath)} bytes)")
    print('='*80)
    try:
        doc = Document(fpath)
        # Print paragraphs
        for i, para in enumerate(doc.paragraphs[:100]):
            if para.text.strip():
                print(f"  [{i:3d}] [{para.style.name:20s}] {para.text[:250]}")
        
        # Check tables
        if doc.tables:
            print(f"\n  --- TABLES: {len(doc.tables)} found ---")
            for ti, table in enumerate(doc.tables[:3]):
                print(f"  Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
                for ri, row in enumerate(table.rows[:5]):
                    cells = [cell.text.strip()[:50] for cell in row.cells]
                    print(f"    Row {ri}: {cells}")
    except Exception as e:
        print(f"  ERROR: {e}")

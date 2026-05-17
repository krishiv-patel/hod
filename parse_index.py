"""
parse_index.py — Build a master lookup of meeting_number → date
from the 3 HOD Index .docx files.
"""
import os
import re
from docx import Document
from datetime import datetime

BASE_DIR = r"d:\HOD\HOD Minutes (Updated on 24 March 2026)\HOD Minutes (Updated on 24 March 2026)"

INDEX_FILES = [
    "HOD Index - 1 to 139.docx",
    "HOD Index - 140 to 288.docx",
    "HOD Index - 289 to .....docx",
]

def parse_meeting_number(text: str) -> int | None:
    """Extract meeting number from strings like '1st', '2nd', '140th', '243rd'."""
    text = text.strip()
    m = re.search(r'(\d+)', text)
    if m:
        return int(m.group(1))
    return None

def parse_date(text: str) -> str | None:
    """
    Parse date from strings like:
      '3 February 2020 | Monday'
      '17 January 2025 | Friday'
      '28 January 2025 | Tuesday | Minutes Not Available'
    Returns ISO date string or None.
    """
    text = text.strip()
    # Remove day-of-week and notes after pipe
    date_part = text.split('|')[0].strip()
    # Also remove any trailing notes
    date_part = re.sub(r'\s*(Minutes Not Available|–).*$', '', date_part, flags=re.IGNORECASE).strip()
    
    # Try various date formats
    for fmt in ['%d %B %Y', '%d %b %Y', '%B %d %Y', '%d-%m-%Y', '%d/%m/%Y']:
        try:
            dt = datetime.strptime(date_part, fmt)
            return dt.strftime('%Y-%m-%d')
        except ValueError:
            continue
    return None

def has_minutes(text: str) -> bool:
    """Check if the date cell indicates minutes are not available."""
    return 'minutes not available' not in text.lower()

def build_meeting_index() -> dict:
    """
    Parse all index files and return a dict:
        { meeting_number (int): { 'date': 'YYYY-MM-DD', 'has_minutes': bool, 'year_seq': str } }
    """
    index = {}
    
    for fname in INDEX_FILES:
        fpath = os.path.join(BASE_DIR, fname)
        if not os.path.exists(fpath):
            print(f"WARNING: Index file not found: {fname}")
            continue
        
        doc = Document(fpath)
        if not doc.tables:
            print(f"WARNING: No tables in {fname}")
            continue
        
        table = doc.tables[0]
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            
            # Skip header rows
            if len(cells) < 3:
                continue
            if cells[0].lower().startswith('sr') or cells[1].lower() == 'date':
                continue
            
            year_seq = cells[0]  # e.g., 'Year 2020 - 1'
            date_str = cells[1]  # e.g., '3 February 2020 | Monday'
            meeting_str = cells[2]  # e.g., '1st'
            
            meeting_num = parse_meeting_number(meeting_str)
            if meeting_num is None:
                continue
            
            parsed_date = parse_date(date_str)
            minutes_available = has_minutes(date_str)
            
            index[meeting_num] = {
                'date': parsed_date,
                'has_minutes': minutes_available,
                'year_seq': year_seq,
            }
    
    return index


if __name__ == '__main__':
    index = build_meeting_index()
    print(f"Total meetings indexed: {len(index)}")
    print(f"Date range: {min(v['date'] for v in index.values() if v['date'])} to {max(v['date'] for v in index.values() if v['date'])}")
    
    # Stats
    with_minutes = sum(1 for v in index.values() if v['has_minutes'])
    without_minutes = sum(1 for v in index.values() if not v['has_minutes'])
    no_date = sum(1 for v in index.values() if v['date'] is None)
    print(f"With minutes: {with_minutes}, Without minutes: {without_minutes}, No date parsed: {no_date}")
    
    # Show some samples
    for num in sorted(index.keys())[:5]:
        print(f"  Meeting {num}: {index[num]}")
    print("  ...")
    for num in sorted(index.keys())[-5:]:
        print(f"  Meeting {num}: {index[num]}")

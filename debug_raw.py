"""Deep inspect Format A files that fell to RAW."""
from docx import Document
import os

BASE_DIR = r"d:\HOD\HOD Minutes (Updated on 24 March 2026)\HOD Minutes (Updated on 24 March 2026)"

# Check a few of the RAW ones
for fname in ["2nd HOD MoM.docx", "36th HOD MoM.docx", "46th HOD MoM.docx", "82nd HOD MoM.docx", "141st HOD MoM.docx", "155th HOD MoM.docx"]:
    fpath = os.path.join(BASE_DIR, fname)
    doc = Document(fpath)
    print(f"\n{'='*80}")
    print(f"FILE: {fname} | Size: {os.path.getsize(fpath)} | Paras: {len(doc.paragraphs)} | Tables: {len(doc.tables)}")
    
    # Show paragraphs 
    non_empty = [(i, p) for i, p in enumerate(doc.paragraphs) if p.text.strip()]
    print(f"Non-empty paragraphs: {len(non_empty)}")
    for i, p in non_empty[:15]:
        print(f"  [{i}] [{p.style.name}] {p.text[:150]}")
    
    # Show tables
    if doc.tables:
        for ti, table in enumerate(doc.tables[:2]):
            print(f"\n  Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
            for ri, row in enumerate(table.rows[:5]):
                cells = [cell.text.strip()[:50] for cell in row.cells]
                print(f"    Row {ri}: {cells}")
    
    # Check if content is in embedded objects
    # Check for OLE or image objects
    rels = doc.part.rels
    print(f"\n  Relationships: {len(rels)} total")
    rel_types = {}
    for r in rels.values():
        rtype = r.reltype.split('/')[-1]
        rel_types[rtype] = rel_types.get(rtype, 0) + 1
    for rtype, count in rel_types.items():
        print(f"    {rtype}: {count}")

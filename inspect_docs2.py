"""Inspect index files and edge cases."""
import os
from docx import Document

BASE_DIR = r"d:\HOD\HOD Minutes (Updated on 24 March 2026)\HOD Minutes (Updated on 24 March 2026)"

# Check the index files
for fname in ["HOD Index - 1 to 139.docx", "HOD Index - 140 to 288.docx", "HOD Index - 289 to .....docx"]:
    fpath = os.path.join(BASE_DIR, fname)
    print(f"\n{'='*80}")
    print(f"INDEX FILE: {fname}")
    print('='*80)
    doc = Document(fpath)
    if doc.tables:
        print(f"  Tables: {len(doc.tables)}")
        for ti, table in enumerate(doc.tables[:2]):
            print(f"  Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
            for ri, row in enumerate(table.rows[:8]):
                cells = [cell.text.strip()[:60] for cell in row.cells]
                print(f"    Row {ri}: {cells}")
            if len(table.rows) > 8:
                print(f"    ... ({len(table.rows) - 8} more rows)")
    else:
        for i, p in enumerate(doc.paragraphs[:20]):
            if p.text.strip():
                print(f"  [{i}] {p.text[:200]}")

# Check the combined file 
print(f"\n{'='*80}")
print("COMBINED FILE: 249 to 270 HOD MOM.docx")
print('='*80)
fpath = os.path.join(BASE_DIR, "249 to 270 HOD MOM.docx")
doc = Document(fpath)
print(f"  Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
for ti, table in enumerate(doc.tables[:3]):
    print(f"  Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
    for ri, row in enumerate(table.rows[:3]):
        cells = [cell.text.strip()[:60] for cell in row.cells]
        print(f"    Row {ri}: {cells}")

# Check 271-273 combined
print(f"\n{'='*80}")
print("COMBINED FILE: 271 272 273 HOD MoM.docx")
print('='*80)
fpath = os.path.join(BASE_DIR, "271 272 273 HOD MoM.docx")
doc = Document(fpath)
print(f"  Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
for ti, table in enumerate(doc.tables[:4]):
    print(f"  Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
    for ri, row in enumerate(table.rows[:3]):
        cells = [cell.text.strip()[:60] for cell in row.cells]
        print(f"    Row {ri}: {cells}")

# Check a mid-range file (table-based, ~460KB)
print(f"\n{'='*80}")
print("MID-RANGE: 163rd HOD MoM.docx")
print('='*80)
fpath = os.path.join(BASE_DIR, "163rd HOD MoM.docx")
doc = Document(fpath)
print(f"  Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
for i, p in enumerate(doc.paragraphs[:10]):
    if p.text.strip():
        print(f"  [{i}] {p.text[:200]}")
if doc.tables:
    for ti, table in enumerate(doc.tables[:2]):
        print(f"  Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
        for ri, row in enumerate(table.rows[:5]):
            cells = [cell.text.strip()[:60] for cell in row.cells]
            print(f"    Row {ri}: {cells}")

# Also check the Draft MoM files 
print(f"\n{'='*80}")
print("DRAFT: Draft MoM of 183rd HOD Meeting.docx")
print('='*80)
fpath = os.path.join(BASE_DIR, "Draft MoM of 183rd HOD Meeting.docx")
doc = Document(fpath)
print(f"  Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
for i, p in enumerate(doc.paragraphs[:10]):
    if p.text.strip():
        print(f"  [{i}] {p.text[:200]}")
if doc.tables:
    t = doc.tables[0]
    print(f"  Table 0: {len(t.rows)} rows x {len(t.columns)} cols")
    for ri, row in enumerate(t.rows[:5]):
        cells = [cell.text.strip()[:80] for cell in row.cells]
        print(f"    Row {ri}: {cells}")

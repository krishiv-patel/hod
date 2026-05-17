"""
parse_meetings.py — Main parser for HOD Meeting Minutes.
Handles 4 document format variants + PDFs.
Outputs meetings_data.csv with structured data.
"""
import os
import re
import sys
import csv
import traceback
from datetime import datetime
from docx import Document
import pandas as pd

# Import the index builder
from parse_index import build_meeting_index

BASE_DIR = r"d:\HOD\HOD Minutes (Updated on 24 March 2026)\HOD Minutes (Updated on 24 March 2026)"
OUTPUT_CSV = r"d:\HOD\meetings_data.csv"

# Files to skip (index files, not meeting minutes)
SKIP_FILES = {
    "HOD Index - 1 to 139.docx",
    "HOD Index - 140 to 288.docx",
    "HOD Index - 289 to .....docx",
}

# ---------- Status Heuristics ----------

STATUS_KEYWORDS = {
    'completed': [
        r'\bcompleted\b', r'\bdone\b', r'\bfinished\b', r'\baccomplished\b',
        r'\bsuccessfully\s+completed\b', r'\bhas been done\b', r'\bwas done\b',
        r'\bsuccessfully\b', r'\bimplemented\b', r'\bexecuted\b',
    ],
    'in_progress': [
        r'\bin progress\b', r'\bongoing\b', r'\bunderway\b', r'\bunder process\b',
        r'\bin process\b', r'\bbeing done\b', r'\bbeing prepared\b',
        r'\bwork is going on\b', r'\bgoing on\b', r'\bunder consideration\b',
        r'\byet to be completed\b', r'\bunder development\b',
    ],
    'pending': [
        r'\bpending\b', r'\byet to be done\b', r'\bnot yet\b', r'\byet to\b',
        r'\bfollow up\b', r'\bfollow-up\b', r'\breminder\b', r'\boverdue\b',
        r'\bnot completed\b', r'\bnot done\b', r'\bmissed\b',
        r'\bstill awaiting\b', r'\bstill pending\b',
    ],
}

def classify_status(text: str) -> str:
    """Classify task status based on keyword heuristics."""
    text_lower = text.lower()
    
    # Check each category
    for status, patterns in STATUS_KEYWORDS.items():
        for pattern in patterns:
            if re.search(pattern, text_lower):
                if status == 'completed':
                    return 'Completed'
                elif status == 'in_progress':
                    return 'In Progress'
                elif status == 'pending':
                    return 'Pending'
    
    # Default: recorded as a resolved/suggested point
    return 'Recorded'


# ---------- Meeting Number Extraction ----------

def extract_meeting_number_from_filename(filename: str) -> list[int]:
    """Extract meeting number(s) from filename."""
    name = os.path.splitext(filename)[0]
    
    # Combined files: "249 to 270 HOD MOM" or "271 272 273 HOD MoM"
    range_match = re.match(r'(\d+)\s+to\s+(\d+)', name)
    if range_match:
        return list(range(int(range_match.group(1)), int(range_match.group(2)) + 1))
    
    multi_match = re.match(r'(\d+)\s+(\d+)\s+(\d+)', name)
    if multi_match:
        return [int(multi_match.group(i)) for i in range(1, 4)]
    
    # Standard: "1st HOD MoM", "134th HOD MoM", "213 HOD MOM", "298 HOD MoM (1st Apr'26)"
    m = re.match(r'(\d+)', name)
    if m:
        return [int(m.group(1))]
    
    # Draft files: "Draft MoM of 183rd HOD Meeting"
    m = re.search(r'(\d+)\s*(?:st|nd|rd|th)?\s*HOD', name, re.IGNORECASE)
    if m:
        return [int(m.group(1))]
    
    return []


def extract_date_from_paragraphs(paragraphs: list) -> str | None:
    """Try to extract date from document header paragraphs."""
    date_patterns = [
        # "Monday, 3 February 2020"
        r'(?:Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday),?\s*(\d{1,2}\s+\w+\s+\d{4})',
        # "3 February 2020"
        r'(\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4})',
        # "Date: 14 November 2022"
        r'Date:\s*(\d{1,2}\s+\w+\s+\d{4})',
        # "24th June 2024"
        r'(\d{1,2}(?:st|nd|rd|th)\s+\w+\s+\d{4})',
    ]
    
    for para in paragraphs[:15]:  # Only check first 15 paragraphs
        text = para.text.strip()
        if not text:
            continue
        for pattern in date_patterns:
            m = re.search(pattern, text, re.IGNORECASE)
            if m:
                date_str = m.group(1)
                # Remove ordinal suffixes
                date_str = re.sub(r'(\d+)(?:st|nd|rd|th)', r'\1', date_str)
                for fmt in ['%d %B %Y', '%d %b %Y']:
                    try:
                        return datetime.strptime(date_str, fmt).strftime('%Y-%m-%d')
                    except ValueError:
                        continue
    return None


# ---------- Format-Specific Parsers ----------

def parse_format_a_paragraphs(doc: Document, meeting_num: int, date: str, filename: str) -> list[dict]:
    """
    Format A: Paragraph-based (early meetings, ~1.9MB files).
    Two sub-variants:
      - Numbered: "1. Football ground to be created... (V. Gole)"
      - Unnumbered: Each paragraph after the header is a separate action item.
    """
    records = []
    paragraphs = doc.paragraphs
    
    # Collect all non-empty paragraph texts
    para_texts = [p.text.strip() for p in paragraphs if p.text.strip()]
    full_text = '\n'.join(para_texts)
    
    # Strategy 1: Try to find numbered items first
    items = re.split(r'\n(?=\d+\.\s)', full_text)
    numbered_records = []
    
    for item in items:
        item = item.strip()
        if not item:
            continue
        m = re.match(r'(\d+)\.\s*(.*)', item, re.DOTALL)
        if not m:
            continue
        task_text = m.group(2).strip()
        assignee_match = re.search(r'\(([^)]+)\)\s*$', task_text)
        assignee = assignee_match.group(1).strip() if assignee_match else 'Not Specified'
        task_text = re.sub(r'\s+', ' ', task_text).strip()
        if len(task_text) < 5:
            continue
        numbered_records.append((task_text, assignee))
    
    if len(numbered_records) >= 3:
        # Use numbered parsing
        for i, (task_text, assignee) in enumerate(numbered_records, 1):
            sr_no = f"{meeting_num}.{i:02d}"
            records.append({
                'meeting_number': meeting_num,
                'meeting_date': date,
                'year': date[:4] if date else '',
                'sr_no': sr_no,
                'task_description': task_text,
                'concerned_unit': assignee,
                'task_status': classify_status(task_text),
                'source_file': filename,
                'format_type': 'A',
            })
        return records
    
    # Strategy 2: Paragraph-per-item (unnumbered)
    # Skip header paragraphs (title, date, "Deliberated..." line)
    start_idx = 0
    skip_keywords = [
        'heads of school', 'hod meeting', 'administrative unit', 'meeting',
        'monday', 'tuesday', 'wednesday', 'thursday', 'friday', 'saturday', 'sunday',
        'deliberated', 'resolved points', 'action by', 'draft minutes', 'final minutes',
        'name of school', 'new plans',
    ]
    
    for i, text in enumerate(para_texts):
        text_lower = text.lower().strip()
        if any(kw in text_lower for kw in skip_keywords):
            start_idx = i + 1
            continue
        # Stop skipping once we hit substantial content
        if len(text) > 30 and start_idx <= i:
            start_idx = i
            break
    
    item_count = 0
    for text in para_texts[start_idx:]:
        text = text.strip()
        if len(text) < 15:
            continue
        
        # Skip meta lines
        text_lower = text.lower()
        if any(kw in text_lower for kw in ['name of school', 'names are as follows', 'new plans']):
            continue
        # Skip lines that are just school names
        if re.match(r'^\d+\.\s*School of', text):
            continue
        if text.startswith('-') and len(text) < 80:
            continue
        
        # Try to extract assignee from parentheses at end
        assignee_match = re.search(r'\(([^)]{2,40})\)\s*$', text)
        assignee = assignee_match.group(1).strip() if assignee_match else 'Not Specified'
        
        task_text = re.sub(r'\s+', ' ', text).strip()
        
        item_count += 1
        sr_no = f"{meeting_num}.{item_count:02d}"
        
        records.append({
            'meeting_number': meeting_num,
            'meeting_date': date,
            'year': date[:4] if date else '',
            'sr_no': sr_no,
            'task_description': task_text,
            'concerned_unit': assignee,
            'task_status': classify_status(task_text),
            'source_file': filename,
            'format_type': 'A',
        })
    
    return records


def parse_format_b_hod_table(doc: Document, meeting_num: int, date: str, filename: str) -> list[dict]:
    """
    Format B: Table with HOD names.
    Columns: [HOD Name/HOD, Discussed Points/Discussed on Points, Remarks/Remarks Made by...]
    Also extracts standalone paragraphs outside the table as additional items.
    """
    records = []
    item_count = 0
    
    # Parse tables
    for table in doc.tables:
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        header_text = ' '.join(headers)
        
        # Determine if this is a HOD-style table
        is_hod_table = any(kw in header_text for kw in [
            'hod', 'name', 'discussed', 'remarks'
        ])
        if not is_hod_table:
            # May be a title table (e.g., meeting title) — skip
            if len(table.rows) <= 2:
                continue
            # Check data rows for sr_no pattern
            first_data = table.rows[1].cells[0].text.strip()
            if re.match(r'\d+\.\d+', first_data):
                # This is actually a Format C table — delegate
                records.extend(parse_format_c_structured_table(
                    table, meeting_num, date, filename
                ))
                continue
        
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) < 2:
                continue
            
            # Skip header rows
            cell0_lower = cells[0].lower()
            if any(h in cell0_lower for h in ['hod', 'name', 'sr', 'sl']):
                continue
            
            hod_name = cells[0]
            discussed = cells[1] if len(cells) > 1 else ''
            remarks = cells[2] if len(cells) > 2 else ''
            
            if not discussed or len(discussed) < 5:
                continue
            
            # Split discussed points by bullet markers, newlines with numbers, or newlines
            points = re.split(r'(?:[·•]\s*|\n\d+\.\s*|\n(?=[A-Z]))', discussed)
            points = [p.strip() for p in points if p.strip() and len(p.strip()) > 5]
            
            # If no split happened, use the whole text
            if not points:
                points = [discussed]
            
            for point in points:
                item_count += 1
                sr_no = f"{meeting_num}.{item_count:02d}"
                
                # Clean the HOD name (take first line, remove prefix dashes)
                hod_clean = hod_name.split('\n')[0].strip()
                hod_clean = re.sub(r'^[-–•·]\s*', '', hod_clean).strip()
                # Remove title prefixes for cleaner grouping
                hod_clean = re.sub(r'^(Dr\.?|Prof\.?|Lt\.? Gen\.?|Col\.?|Brig\.?)\s*', r'\1 ', hod_clean).strip()
                
                combined_text = point
                if remarks and len(remarks) > 5:
                    combined_text += f" [Remarks: {remarks[:300]}]"
                
                records.append({
                    'meeting_number': meeting_num,
                    'meeting_date': date,
                    'year': date[:4] if date else '',
                    'sr_no': sr_no,
                    'task_description': re.sub(r'\s+', ' ', combined_text).strip(),
                    'concerned_unit': hod_clean,
                    'task_status': classify_status(combined_text),
                    'source_file': filename,
                    'format_type': 'B',
                })
    
    # Also parse standalone paragraphs outside tables (resolved points in body)
    para_records = _extract_standalone_paragraphs(doc, meeting_num, date, filename, item_count)
    records.extend(para_records)
    
    return records


def _extract_standalone_paragraphs(doc: Document, meeting_num: int, date: str, filename: str, start_count: int) -> list[dict]:
    """Extract resolved/suggested points from paragraphs outside tables."""
    records = []
    item_count = start_count
    in_content = False
    
    skip_keywords = [
        'final minutes', 'draft minutes', 'hod meeting', 'meeting',
        'deliberated', 'resolved points', 'action by', '---', '—--',
        'staff/unit', 'the hod meeting',
    ]
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or len(text) < 20:
            continue
        
        text_lower = text.lower()
        
        # Detect content start
        if any(kw in text_lower for kw in ['deliberated', 'resolved points', 'action by']):
            in_content = True
            continue
        
        if not in_content:
            # Also start content if we see date/time patterns pass
            if re.match(r'\d{1,2}:\d{2}\s*(AM|PM)', text, re.IGNORECASE):
                in_content = True
                continue
            continue
        
        # Skip meta lines
        if any(kw in text_lower for kw in skip_keywords):
            continue
        if text.startswith('---') or text.startswith('—'):
            continue
        
        # This is a content paragraph — extract as a task
        # Check if it starts with "The HOD" pattern (common in these docs)
        if len(text) >= 30:
            item_count += 1
            sr_no = f"{meeting_num}.{item_count:02d}"
            
            # Try to extract assignee
            assignee = 'Not Specified'
            # Look for patterns like "(Director-SPICSM)" or just named entities
            assignee_match = re.search(r'\(([^)]{2,60})\)\s*$', text)
            if assignee_match:
                assignee = assignee_match.group(1).strip()
            
            records.append({
                'meeting_number': meeting_num,
                'meeting_date': date,
                'year': date[:4] if date else '',
                'sr_no': sr_no,
                'task_description': re.sub(r'\s+', ' ', text).strip(),
                'concerned_unit': assignee,
                'task_status': classify_status(text),
                'source_file': filename,
                'format_type': 'B',
            })
    
    return records


def parse_format_c_structured_table(table, meeting_num: int, date: str, filename: str) -> list[dict]:
    """
    Format C: Structured table (modern format).
    Columns: [Sr No, Resolved/Suggested Points, Concerned Unit/Staff]
    """
    records = []
    
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 2:
            continue
        
        sr_no_text = cells[0]
        
        # Skip header rows
        if any(h in sr_no_text.lower() for h in ['sr', 'no', 'sl', 'serial', 'hod name']):
            continue
        
        # Extract meeting number from sr_no (e.g., "163.01" → 163)
        sr_match = re.match(r'(\d+)[.\s]*(\d+)?', sr_no_text)
        if not sr_match:
            continue
        
        doc_meeting_num = int(sr_match.group(1))
        item_num = sr_match.group(2) or '0'
        sr_no = f"{doc_meeting_num}.{int(item_num):02d}"
        
        task_text = cells[1] if len(cells) > 1 else ''
        concerned = cells[2] if len(cells) > 2 else 'Not Specified'
        
        if not task_text or len(task_text) < 5:
            continue
        
        records.append({
            'meeting_number': doc_meeting_num,
            'meeting_date': date,
            'year': date[:4] if date else '',
            'sr_no': sr_no,
            'task_description': re.sub(r'\s+', ' ', task_text).strip(),
            'concerned_unit': re.sub(r'\s+', ' ', concerned).strip(),
            'task_status': classify_status(task_text),
            'source_file': filename,
            'format_type': 'C',
        })
    
    return records


def detect_format(doc: Document, filename: str, file_size: int) -> str:
    """Detect document format based on structure analysis."""
    has_tables = len(doc.tables) > 0
    
    # Combined files with multiple meetings
    if re.match(r'\d+\s+to\s+\d+', filename) or re.match(r'\d+\s+\d+\s+\d+', filename):
        return 'D'
    
    if not has_tables:
        return 'A'
    
    # Check ALL tables to find the data table (some files have a title table first)
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        header_text = ' '.join(headers)
        
        # Format B: HOD-name based discussions
        if any(kw in header_text for kw in [
            'hod name', 'name of the hod', 'discussed on points',
            'discussed points', 'remarks made by',
        ]):
            return 'B'
        
        # Also catch single-word 'hod' header in first cell
        if headers and headers[0].strip() == 'hod' and len(headers) >= 2:
            if 'discussed' in header_text or 'remarks' in header_text:
                return 'B'
        
        # Format C: Sr No based structured table
        if any(x in header_text for x in ['sr no', 'sr.', 'resolved', 'suggested', 'sl no']):
            return 'C'
        
        # Check data rows for sr_no pattern (e.g., "163.01")
        for row in table.rows[1:3]:
            first_cell = row.cells[0].text.strip()
            if re.match(r'\d+\.\d+', first_cell):
                return 'C'
    
    # Large files (~1.9MB) without matching headers — paragraph-based
    if file_size > 1_500_000:
        return 'A'
    
    # Small files with tables we couldn't classify — try B then C
    for table in doc.tables:
        if len(table.rows) >= 3:
            return 'B'
    
    return 'A'  # Fallback


def parse_pdf_file(filepath: str, meeting_nums: list[int], meeting_index: dict) -> list[dict]:
    """Parse PDF files using pdfplumber."""
    records = []
    filename = os.path.basename(filepath)
    
    try:
        import pdfplumber
        
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                # Try to extract tables first
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        for row in table:
                            if not row or len(row) < 2:
                                continue
                            
                            sr_no_text = (row[0] or '').strip()
                            sr_match = re.match(r'(\d+)[.\s]*(\d+)?', sr_no_text)
                            if not sr_match:
                                continue
                            
                            # Skip headers
                            if any(h in sr_no_text.lower() for h in ['sr', 'no', 'sl']):
                                continue
                            
                            doc_meeting_num = int(sr_match.group(1))
                            item_num = sr_match.group(2) or '0'
                            sr_no = f"{doc_meeting_num}.{int(item_num):02d}"
                            
                            task_text = (row[1] or '').strip() if len(row) > 1 else ''
                            concerned = (row[2] or '').strip() if len(row) > 2 else 'Not Specified'
                            
                            if not task_text or len(task_text) < 5:
                                continue
                            
                            # Get date from index
                            idx_data = meeting_index.get(doc_meeting_num, {})
                            date = idx_data.get('date', '')
                            
                            records.append({
                                'meeting_number': doc_meeting_num,
                                'meeting_date': date,
                                'year': date[:4] if date else '',
                                'sr_no': sr_no,
                                'task_description': re.sub(r'\s+', ' ', task_text).strip(),
                                'concerned_unit': re.sub(r'\s+', ' ', concerned).strip(),
                                'task_status': classify_status(task_text),
                                'source_file': filename,
                                'format_type': 'PDF',
                            })
                else:
                    # Fallback: extract text and parse like Format A
                    text = page.extract_text()
                    if text:
                        for meeting_num in meeting_nums:
                            idx_data = meeting_index.get(meeting_num, {})
                            date = idx_data.get('date', '')
                            
                            items = re.split(r'\n(?=\d+\.\s)', text)
                            item_count = 0
                            for item in items:
                                m = re.match(r'(\d+)\.\s*(.*)', item, re.DOTALL)
                                if not m:
                                    continue
                                task_text = m.group(2).strip()
                                if len(task_text) < 5:
                                    continue
                                item_count += 1
                                sr_no = f"{meeting_num}.{item_count:02d}"
                                records.append({
                                    'meeting_number': meeting_num,
                                    'meeting_date': date,
                                    'year': date[:4] if date else '',
                                    'sr_no': sr_no,
                                    'task_description': re.sub(r'\s+', ' ', task_text).strip(),
                                    'concerned_unit': 'Not Specified',
                                    'task_status': classify_status(task_text),
                                    'source_file': filename,
                                    'format_type': 'PDF',
                                })
    except ImportError:
        print("WARNING: pdfplumber not installed. Skipping PDF files.")
    except Exception as e:
        print(f"ERROR parsing PDF {filename}: {e}")
    
    return records


def parse_single_docx(filepath: str, meeting_index: dict) -> list[dict]:
    """Parse a single .docx file and return list of record dicts."""
    filename = os.path.basename(filepath)
    file_size = os.path.getsize(filepath)
    
    try:
        doc = Document(filepath)
    except Exception as e:
        print(f"ERROR: Cannot read {filename}: {e}")
        return []
    
    # Get meeting number(s) from filename
    meeting_nums = extract_meeting_number_from_filename(filename)
    if not meeting_nums:
        print(f"WARNING: Cannot extract meeting number from {filename}")
        return []
    
    primary_meeting_num = meeting_nums[0]
    
    # Get date from index, falling back to document header
    idx_data = meeting_index.get(primary_meeting_num, {})
    date = idx_data.get('date', '')
    if not date:
        date = extract_date_from_paragraphs(doc.paragraphs) or ''
    
    # Detect format
    fmt = detect_format(doc, filename, file_size)
    
    records = []
    
    if fmt == 'A':
        records = parse_format_a_paragraphs(doc, primary_meeting_num, date, filename)
    
    elif fmt == 'B':
        records = parse_format_b_hod_table(doc, primary_meeting_num, date, filename)
    
    elif fmt == 'C':
        if doc.tables:
            for table in doc.tables:
                records.extend(parse_format_c_structured_table(
                    table, primary_meeting_num, date, filename
                ))
    
    elif fmt == 'D':
        # Combined multi-meeting file — each table is a separate meeting
        if doc.tables:
            for table in doc.tables:
                # The meeting number is embedded in the Sr No column
                table_records = parse_format_c_structured_table(
                    table, primary_meeting_num, date, filename
                )
                # Update dates from index for each record's actual meeting number
                for rec in table_records:
                    actual_num = rec['meeting_number']
                    actual_data = meeting_index.get(actual_num, {})
                    actual_date = actual_data.get('date', '')
                    if actual_date:
                        rec['meeting_date'] = actual_date
                        rec['year'] = actual_date[:4]
                records.extend(table_records)
    
    if not records:
        # Last resort: try to extract any text as a single record
        all_text = ' '.join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        if len(all_text) > 20:
            records.append({
                'meeting_number': primary_meeting_num,
                'meeting_date': date,
                'year': date[:4] if date else '',
                'sr_no': f"{primary_meeting_num}.00",
                'task_description': all_text[:2000],
                'concerned_unit': 'Not Specified',
                'task_status': 'Recorded',
                'source_file': filename,
                'format_type': 'RAW',
            })
    
    return records


def main():
    print("=" * 60)
    print("HOD Meeting Minutes Parser")
    print("=" * 60)
    
    # Step 1: Build index
    print("\n[1/3] Building meeting index...")
    meeting_index = build_meeting_index()
    print(f"      Indexed {len(meeting_index)} meetings")
    
    # Step 2: Collect all files
    print("\n[2/3] Scanning files...")
    all_files = []
    for fname in os.listdir(BASE_DIR):
        if fname in SKIP_FILES:
            continue
        fpath = os.path.join(BASE_DIR, fname)
        if os.path.isdir(fpath):
            continue
        all_files.append((fname, fpath))
    
    docx_files = [(f, p) for f, p in all_files if f.lower().endswith('.docx')]
    pdf_files = [(f, p) for f, p in all_files if f.lower().endswith('.pdf')]
    print(f"      Found {len(docx_files)} .docx files and {len(pdf_files)} .pdf files")
    
    # Step 3: Parse all files
    print("\n[3/3] Parsing documents...")
    all_records = []
    format_counts = {'A': 0, 'B': 0, 'C': 0, 'D': 0, 'PDF': 0, 'RAW': 0}
    error_files = []
    
    for i, (fname, fpath) in enumerate(docx_files):
        try:
            records = parse_single_docx(fpath, meeting_index)
            all_records.extend(records)
            for rec in records:
                format_counts[rec['format_type']] = format_counts.get(rec['format_type'], 0) + 1
            if (i + 1) % 25 == 0:
                print(f"      Processed {i + 1}/{len(docx_files)} files ({len(all_records)} records so far)")
        except Exception as e:
            error_files.append((fname, str(e)))
            print(f"      ERROR: {fname}: {e}")
            traceback.print_exc()
    
    # Parse PDFs
    for fname, fpath in pdf_files:
        try:
            meeting_nums = extract_meeting_number_from_filename(fname)
            records = parse_pdf_file(fpath, meeting_nums, meeting_index)
            all_records.extend(records)
            format_counts['PDF'] += len(records)
            print(f"      PDF: {fname} → {len(records)} records")
        except Exception as e:
            error_files.append((fname, str(e)))
            print(f"      ERROR PDF: {fname}: {e}")
    
    # Step 4: Save to CSV
    print(f"\nSaving {len(all_records)} records to CSV...")
    df = pd.DataFrame(all_records)
    
    # Sort by meeting number and sr_no
    df['_sort_key'] = df['meeting_number'] * 1000 + df['sr_no'].apply(
        lambda x: float(x.split('.')[-1]) if '.' in str(x) else 0
    )
    df = df.sort_values('_sort_key').drop(columns=['_sort_key'])
    df.to_csv(OUTPUT_CSV, index=False, encoding='utf-8-sig')
    
    # Summary
    print("\n" + "=" * 60)
    print("PARSING COMPLETE")
    print("=" * 60)
    print(f"  Total records extracted:   {len(all_records)}")
    print(f"  Unique meetings:           {df['meeting_number'].nunique()}")
    print(f"  Date range:                {df[df['meeting_date'] != '']['meeting_date'].min()} to {df[df['meeting_date'] != '']['meeting_date'].max()}")
    print(f"\n  Format breakdown:")
    for fmt, count in sorted(format_counts.items()):
        if count > 0:
            print(f"    {fmt}: {count} records")
    
    print(f"\n  Status breakdown:")
    for status, count in df['task_status'].value_counts().items():
        print(f"    {status}: {count}")
    
    print(f"\n  Top 10 departments/units:")
    for unit, count in df['concerned_unit'].value_counts().head(10).items():
        print(f"    {unit[:50]:50s} | {count} tasks")
    
    if error_files:
        print(f"\n  ⚠ Errors in {len(error_files)} files:")
        for fname, err in error_files:
            print(f"    {fname}: {err}")
    
    print(f"\n  Output saved to: {OUTPUT_CSV}")
    return df


if __name__ == '__main__':
    main()
